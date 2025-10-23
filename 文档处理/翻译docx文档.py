from translate import Translator
from docx import Document

def translate_docx(doc_path, target_language, output_path):
    translator = Translator(to_lang=target_language)
    doc = Document(doc_path)
    total_paragraphs = len(doc.paragraphs)
    translated_count = 0
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() != "":
            # 翻译段落文本
            translated_text = translator.translate(paragraph.text)
            paragraph.text = translated_text
            translated_count += 1
            print(f'已完成 {translated_count}/{total_paragraphs} 段落翻译')
    
    # 保存翻译后的文档
    doc.save(output_path)

# 调用函数进行翻译，假设目标语言是中文
translate_docx('abc.docx', 'zh', 'abc_trans.docx')
