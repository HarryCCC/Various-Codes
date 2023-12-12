from docx import Document

def extract_pages(doc_path, output_path, pages_to_extract):
    # 加载文档
    doc = Document(doc_path)
    new_doc = Document()
    
    # 计算每页的段落数（假设每页大约有多少个段落）
    paragraphs_per_page = 50  # 这个值可能需要根据实际文档调整
    
    # 提取前 10 页的段落
    for i in range(pages_to_extract * paragraphs_per_page):
        if i < len(doc.paragraphs):
            paragraph = doc.paragraphs[i]
            new_doc.add_paragraph(paragraph.text, style=paragraph.style)
    
    # 保存新文档
    new_doc.save(output_path)

# 调用函数提取前 10 页
extract_pages('abc.docx', 'abc_extracted.docx', 10)
